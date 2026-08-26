"""Pull the plain text out of an uploaded CV.

Only PDF and DOCX are accepted. Anything else, or a file we cannot read,
raises CVExtractionError with a message that is safe to show the user.
"""

import io
import re
import zipfile

import pdfplumber
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

MAX_FILE_BYTES = 5 * 1024 * 1024  # 5 MB

PDF_CONTENT_TYPES = {"application/pdf", "application/x-pdf"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}

# A CV with less text than this is almost certainly a scan or an empty file
MIN_USEFUL_CHARS = 30

WHITESPACE_RE = re.compile(r"[ \t ]+")
BLANK_LINES_RE = re.compile(r"\n{3,}")


class CVExtractionError(Exception):
    """Raised when a file cannot be accepted or read."""


def _clean(text: str) -> str:
    """Tidy the extracted text without destroying its structure."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ")
    # Collapse runs of spaces and tabs but keep the line breaks
    text = WHITESPACE_RE.sub(" ", text)
    text = re.sub(r" *\n *", "\n", text)
    # At most one blank line between blocks
    text = BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


def detect_kind(filename: str | None, content_type: str | None) -> str:
    """Work out whether this is a pdf or a docx.

    The extension is checked first because browsers are inconsistent about the
    content type they send, then the content type is used as a fallback.
    """
    name = (filename or "").lower()
    ctype = (content_type or "").lower().split(";")[0].strip()

    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"

    if ctype in PDF_CONTENT_TYPES:
        return "pdf"
    if ctype in DOCX_CONTENT_TYPES:
        return "docx"

    if name.endswith(".doc"):
        raise CVExtractionError(
            "Old .doc files are not supported. Please save the CV as PDF or "
            ".docx and upload it again."
        )

    raise CVExtractionError(
        "Unsupported file type. Please upload your CV as a PDF or a DOCX file."
    )


def _extract_pdf(data: bytes) -> str:
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            if not pdf.pages:
                raise CVExtractionError("This PDF has no pages.")
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n\n".join(pages)
    except CVExtractionError:
        raise
    except Exception as exc:
        raise CVExtractionError(
            "This PDF could not be read. It may be corrupted or password "
            "protected."
        ) from exc


def _extract_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        raise CVExtractionError(
            "This DOCX file could not be read. It may be corrupted, or it may "
            "be an old .doc file renamed to .docx."
        ) from exc
    except Exception as exc:
        raise CVExtractionError("This DOCX file could not be read.") from exc

    parts = [p.text for p in document.paragraphs]

    # Plenty of CVs lay everything out inside tables, so read those too
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    return "\n".join(parts)


def extract_text(
    data: bytes, filename: str | None = None, content_type: str | None = None
) -> tuple[str, str]:
    """Extract clean text from CV file bytes.

    Returns the detected kind ("pdf" or "docx") and the cleaned text.
    Raises CVExtractionError for anything the user needs to fix.
    """
    if not data:
        raise CVExtractionError("The uploaded file is empty.")

    if len(data) > MAX_FILE_BYTES:
        limit_mb = MAX_FILE_BYTES // (1024 * 1024)
        raise CVExtractionError(
            f"File is too large. The maximum size is {limit_mb} MB."
        )

    kind = detect_kind(filename, content_type)
    raw = _extract_pdf(data) if kind == "pdf" else _extract_docx(data)
    text = _clean(raw)

    if len(text) < MIN_USEFUL_CHARS:
        raise CVExtractionError(
            "No readable text was found in this file. If your CV is a scanned "
            "image, please upload a version with selectable text."
        )

    return kind, text
